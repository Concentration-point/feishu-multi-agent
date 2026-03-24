from __future__ import annotations

from pathlib import Path

from docx import Document

from .common import ensure_dir


def docx_build(directory: str) -> dict:
    root = Path(directory).resolve()
    transcripts_dir = root / "transcripts"
    docs_dir = ensure_dir(root / "docs")

    items = []
    transcript_files = sorted(transcripts_dir.glob("*.md")) if transcripts_dir.exists() else []

    for transcript in transcript_files:
        doc = Document()
        doc.add_heading(transcript.stem, level=1)
        content = transcript.read_text(encoding="utf-8")
        for block in content.split("\n\n"):
            doc.add_paragraph(block.strip())
        out_path = docs_dir / f"{transcript.stem}.docx"
        doc.save(out_path)
        items.append({"source": str(transcript), "docx": str(out_path), "status": "ok"})

    status = "ok" if items else "empty"
    return {
        "status": status,
        "directory": str(root),
        "items": items,
        "message": "Word 文档生成完了。" if items else "还没找到 transcript，先别急着出 Word。",
    }
